import tkinter as tk
from tkinter import ttk, messagebox
from tkinter.filedialog import asksaveasfilename
import csv
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from pypinyin import lazy_pinyin
from tkinter.filedialog import askopenfilename
from openpyxl import load_workbook
import os

def load_name_map(filepath="name_id_map.csv"):
    name_map = {}

    # 如果文件不存在，自动创建一个空模板
    if not os.path.exists(filepath):
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("中文名,英文名,ID\n")
        print(" 自动创建空的 name_id_map.csv 文件")

    try:
        with open(filepath, encoding='utf-8') as f:
            reader = csv.DictReader(f)
            if not reader.fieldnames or '中文名' not in reader.fieldnames or 'ID' not in reader.fieldnames:
                messagebox.showwarning("警告", "name_id_map.csv 文件格式无效，将跳过自动填充")
                return name_map

            has_eng = '英文名' in reader.fieldnames
            for row in reader:
                name = row['中文名']
                id_ = row['ID']
                if not name or not id_:
                    continue
                engname = row['英文名'] if has_eng and row.get('英文名') else ' '.join(p.upper() for p in lazy_pinyin(name))
                name_map[name] = {'英文名': engname, 'ID': id_}

    except Exception as e:
        messagebox.showerror("错误", f"读取 name_id_map.csv 失败：{e}")

    return name_map


# 读取技能章翻译表
def load_badge_translation(filepath="badge_translation.csv"):
    badge_map = {}
    try:
        with open(filepath, encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            print(" 字段名列表：", reader.fieldnames)

            chinese_key = '中文技能章'
            korean_key = '韩文技能章'
            if '\ufeff中文技能章' in reader.fieldnames:
                chinese_key = '\ufeff中文技能章'

            count = 0
            for row in reader:
                print(f"➡️ 读取行：{row}")
                if chinese_key in row and korean_key in row:
                    badge_map[row[chinese_key]] = row[korean_key]
                    count += 1

            print(f" 最终读取技能章数量：{count}")
    except Exception as e:
        messagebox.showerror("错误", f"读取失败：{e}")
    return badge_map

# 读取姓名-ID映射表
def load_name_map(filepath="name_id_map.csv"):
    name_map = {}
    if not os.path.exists(filepath):
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("中文名,英文名,ID\n")
        print("✅ 自动创建空的 name_id_map.csv 文件")

    try:
        with open(filepath, encoding='utf-8') as f:
            reader = csv.DictReader(f)

            if reader.fieldnames is None:
                messagebox.showwarning("警告", "name_id_map.csv 是空文件，将跳过自动填充")
                return name_map

            if '中文名' not in reader.fieldnames or 'ID' not in reader.fieldnames:
                messagebox.showwarning("警告", "name_id_map.csv 缺少必要的列（中文名、ID）")
                return name_map

            has_eng = '英文名' in reader.fieldnames
            for row in reader:
                name = row['中文名']
                id_ = row['ID']
                if not name or not id_:
                    continue
                engname = row['英文名'] if has_eng and row.get('英文名') else ' '.join(p.upper() for p in lazy_pinyin(name))
                name_map[name] = {'英文名': engname, 'ID': id_}

    except Exception as e:
        messagebox.showerror("错误", f"读取 name_id_map.csv 失败：{e}")

    return name_map

# 主应用类
class Application:
    def __init__(self, root):
        self.root = root
        self.root.title("技能章申请书生成器")
        self.name_map = load_name_map()
        self.badge_translation = load_badge_translation()
        self.entries = []
        self.create_widgets()

    def create_widgets(self):
        frame = ttk.Frame(self.root, padding=10)
        frame.grid(row=0, column=0)
    


        self.level_var = tk.StringVar()
        self.badge_var = tk.StringVar()
        self.grade_var = tk.StringVar()
        self.date_var = tk.StringVar()
        self.group_var = tk.StringVar()
        self.leader_var = tk.StringVar()
        self.name_var = tk.StringVar()
        self.engname_var = tk.StringVar()
        self.id_var = tk.StringVar()
        self.note_var = tk.StringVar()

        # 表格预览区域
        self.tree = ttk.Treeview(frame, columns=[str(i) for i in range(1, 10)], show='headings', height=8)
        headers = ["序号", "姓名", "英文名", "ID", "现级别", "技能章名", "承认日", "年级", "备注"]
        for i, header in enumerate(headers):
            self.tree.heading(str(i+1), text=header)
            self.tree.column(str(i+1), width=80, anchor='center')
        self.tree.grid(row=8, column=0, columnspan=4, pady=(10, 0))
        self.tree.bind("<Double-1>", self.edit_selected_row)



        # 团名与队长姓名（最上面）
        ttk.Label(frame, text="团名:").grid(row=0, column=0)
        ttk.Entry(frame, textvariable=self.group_var).grid(row=0, column=1)

        ttk.Label(frame, text="队长姓名:").grid(row=0, column=2)
        ttk.Entry(frame, textvariable=self.leader_var).grid(row=0, column=3)

        #  现级别与技能章
        ttk.Label(frame, text="现级别(현급위):").grid(row=1, column=0)
        ttk.Combobox(frame, textvariable=self.level_var, values=["初级(초급)", "二级(2급)", "一级(1급)", "星级(별급)", "国花(무궁화)"]).grid(row=1, column=1)

        ttk.Label(frame, text="技能章(기능장명):").grid(row=1, column=2)
        self.badge_names = list(self.badge_translation.keys())
        self.badge_entry = ttk.Entry(frame, textvariable=self.badge_var)
        self.badge_entry.grid(row=1, column=3)
        self.badge_entry.bind('<KeyRelease>', self.filter_badges)

        self.badge_listbox = tk.Listbox(frame, height=5)
        self.badge_listbox.grid(row=2, column=3, sticky="we")
        self.badge_listbox.bind("<<ListboxSelect>>", self.select_badge_from_listbox)
        self.update_badge_listbox()

        #  年级与承认日
        ttk.Label(frame, text="年级(구분):").grid(row=3, column=0)
        ttk.Combobox(frame, textvariable=self.grade_var, values=["小学", "初中", "高中"]).grid(row=3, column=1)

        ttk.Label(frame, text="承认日(인가일) (YYYY-MM-DD):").grid(row=3, column=2)
        ttk.Entry(frame, textvariable=self.date_var).grid(row=3, column=3)

        #  中文名与英文名生成按钮
        ttk.Label(frame, text="中文姓名:").grid(row=4, column=0)
        ttk.Entry(frame, textvariable=self.name_var).grid(row=4, column=1)
        ttk.Button(frame, text="自动生成英文名", command=self.generate_english_name).grid(row=4, column=2)

        #  英文名与 ID
        ttk.Label(frame, text="英文名:").grid(row=5, column=0)
        ttk.Label(frame, textvariable=self.engname_var).grid(row=5, column=1)
        ttk.Label(frame, text="ID:").grid(row=5, column=2)
        ttk.Entry(frame, textvariable=self.id_var).grid(row=5, column=3)

        #  备注
        ttk.Label(frame, text="备注:").grid(row=6, column=0)
        ttk.Entry(frame, textvariable=self.note_var).grid(row=6, column=1, columnspan=3, sticky="we")

        #  操作按钮
        ttk.Button(frame, text="添加到列表", command=self.add_to_list).grid(row=7, column=0, columnspan=2)
        ttk.Button(frame, text="生成 DOCX", command=self.generate_docx).grid(row=7, column=2, columnspan=2)

        ttk.Button(frame, text="导入姓名-ID Excel", command=self.import_excel_name_map).grid(row=9, column=0, columnspan=4, pady=(10, 0))

        ttk.Button(frame, text="删除选中行", command=self.delete_selected_row).grid(row=10, column=0, columnspan=4, pady=(5, 0))



    def update_badge_listbox(self, filtered=None):
        self.badge_listbox.delete(0, tk.END)
        for badge in filtered or self.badge_names:
            self.badge_listbox.insert(tk.END, badge)

    def filter_badges(self, event):
        typed = self.badge_var.get()
        filtered = [b for b in self.badge_names if typed in b]
        self.update_badge_listbox(filtered)

    def select_badge_from_listbox(self, event):
        selection = self.badge_listbox.curselection()
        if selection:
            self.badge_var.set(self.badge_listbox.get(selection[0]))
            self.badge_listbox.selection_clear(0, tk.END)

    def generate_english_name(self):
        name = self.name_var.get()
        if name:
            pinyin = lazy_pinyin(name)
            self.engname_var.set(' '.join(p.upper() for p in pinyin))

    def add_to_list(self):
        raw_date = self.date_var.get()
        formatted_date = raw_date

        if raw_date.isdigit() and len(raw_date) == 8:
            try:
                y, m, d = raw_date[:4], raw_date[4:6], raw_date[6:8]
                datetime.strptime(f"{y}-{m}-{d}", "%Y-%m-%d")  # 校验
                formatted_date = f"{y}-{m}-{d}"
            except ValueError:
                pass  # 无效格式跳过转换

        data = [
            len(self.entries) + 1,
            self.name_var.get(),
            self.engname_var.get(),
            self.id_var.get(),
            self.level_var.get(),
            self.badge_var.get(),
            formatted_date,
            self.grade_var.get(),
            self.note_var.get()
        ]

        self.entries.append(data)
        self.tree.insert('', tk.END, values=data)
      


   

    def generate_docx(self):
        if not self.entries:
            messagebox.showwarning("无数据", "请先添加人员信息")
            return

        doc = Document()
        doc.add_heading("기능장 인가 보고서", 0)
        table = doc.add_table(rows=1, cols=9)
        headers = ["순", "성명姓名", "영어면英文名", "ID", "현급위现级别", "기능장명技能章名", "承认日", "S.V 구분年级", "备注비  고"]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for row in self.entries:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                if i == 5:
                    korean_badge = self.badge_translation.get(val, val)
                    cells[i].text = korean_badge
                else:
                    cells[i].text = str(val)

        # 使用最后一条数据的承认日作为申请日期
        apply_date = self.entries[-1][6]
        try:
            if len(apply_date) == 8 and apply_date.isdigit():
                y = int(apply_date[:4])
                m = int(apply_date[4:6])
                d = int(apply_date[6:8])
                date_str = f"     {y}년 {m}월 {d}일"
            else:
                raise ValueError
        except:
            date_str = "     ????년 ??월 ??일"


        doc.add_paragraph("\n\n")
        para = doc.add_paragraph("위와 같이 기능장 취득을 인가하고 보고합니다.")
        para.alignment = 1

        doc.add_paragraph(date_str).alignment = 1

        group_name = self.group_var.get() or "???"
        leader_name = self.leader_var.get() or "???"

        line = f"제   {group_name}    团대(단)         대장队장    {leader_name}      (인)"
        doc.add_paragraph("\n" + line).alignment = 1

        doc.add_paragraph("\n한국스카우트 서울북부연맹  귀하").alignment = 1

        filename = asksaveasfilename(defaultextension=".docx", filetypes=[("Word 文档", "*.docx")])
        if filename:
            doc.save(filename)
            messagebox.showinfo("成功", f"已保存到 {filename}")

    def delete_selected_row(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("未选择", "请先选中一行再删除")
            return

        confirm = messagebox.askyesno("确认删除", "确定要删除选中的这条记录吗？")
        if not confirm:
            return

        index = self.tree.index(selected_item)
        self.tree.delete(selected_item)
        del self.entries[index]

        # 重新编号
        for i, item in enumerate(self.tree.get_children()):
            values = list(self.tree.item(item)['values'])
            values[0] = i + 1
            self.tree.item(item, values=values)
            self.entries[i][0] = i + 1
 

    def edit_selected_row(self, event):
        selected_item = self.tree.focus()
        if not selected_item:
            return

        values = self.tree.item(selected_item)['values']
        if not values:
            return

        edit_win = tk.Toplevel(self.root)
        edit_win.title("编辑数据")

        labels = ["姓名", "英文名", "ID", "现级别", "技能章名", "承认日", "年级", "备注"]
        vars_ = []

        for i, (label, val) in enumerate(zip(labels, values[1:])):
            tk.Label(edit_win, text=label + ":").grid(row=i, column=0, padx=5, pady=3, sticky='e')
            var = tk.StringVar(value=val)
            tk.Entry(edit_win, textvariable=var).grid(row=i, column=1, padx=5, pady=3, sticky='w')
            vars_.append(var)

        def save_changes():
            edited = [v.get() for v in vars_]

            # ✅ 自动转换日期格式（第6列，对应承认日）
            raw_date = edited[5]  # 第6项是“承认日”
            if raw_date.isdigit() and len(raw_date) == 8:
                try:
                    y, m, d = raw_date[:4], raw_date[4:6], raw_date[6:8]
                    datetime.strptime(f"{y}-{m}-{d}", "%Y-%m-%d")  # 验证是否为有效日期
                    edited[5] = f"{y}-{m}-{d}"
                except ValueError:
                    pass  # 如果不是合法日期就不处理

            new_values = [values[0]] + edited
            self.tree.item(selected_item, values=new_values)
            self.entries[int(values[0]) - 1] = new_values
            edit_win.destroy()



        tk.Button(edit_win, text="保存修改", command=save_changes).grid(row=len(labels), column=0, columnspan=2, pady=10)


    def import_excel_name_map(self):
        filepath = askopenfilename(filetypes=[("Excel 文件", "*.xlsx")])
        if not filepath:
            return

        try:
            wb = load_workbook(filepath)
            sheet = wb.active

            headers = [cell.value for cell in sheet[1]]
            col_map = {h: i for i, h in enumerate(headers)}

            if '中文名' not in col_map or 'ID' not in col_map:
                messagebox.showerror("错误", "Excel 文件必须包含 '中文名' 和 'ID' 列")
                return

            has_eng = '英文名' in col_map
            imported = 0

            for row in sheet.iter_rows(min_row=2, values_only=True):
                name = row[col_map['中文名']]
                id_ = str(row[col_map['ID']]) if row[col_map['ID']] else ''
                if not name or not id_:
                    continue
                if has_eng and row[col_map['英文名']]:
                    engname = row[col_map['英文名']]
                else:
                    engname = ' '.join(p.upper() for p in lazy_pinyin(name))

                self.name_map[name] = {'英文名': engname, 'ID': id_}
                imported += 1

            messagebox.showinfo("导入成功", f"成功导入 {imported} 条记录")

        except Exception as e:
            messagebox.showerror("错误", f"导入失败：{e}")


# 启动程序
if __name__ == '__main__':
    root = tk.Tk()
    app = Application(root)
    root.mainloop()
